import pandas as pd
import docx
import os

cwd = r"f:\Антигравити ии\калькулятор материалов"

print("--- Ткань ---")
df_tkan = pd.read_excel(os.path.join(cwd, "БД ткани для ноутбук ЛМ.xlsx"))
print(df_tkan.columns.tolist())
print(df_tkan.head(3))

print("\n--- Поролон ---")
df_porolon = pd.read_excel(os.path.join(cwd, "цены поролон (новые) Moncor.xlsx"))
print(df_porolon.columns.tolist())
print(df_porolon.head(3))

print("\n--- Фанера ---")
doc_fanera = docx.Document(os.path.join(cwd, "Прайс-лист на фанеру (Латвия).docx"))
for t in doc_fanera.tables:
    for row in t.rows[:3]:
        print([cell.text for cell in row.cells])

print("\n--- МДФ/ДСП ---")
doc_mdf = docx.Document(os.path.join(cwd, "Цены на МДФ и ДСП.docx"))
for t in doc_mdf.tables:
    for row in t.rows[:3]:
        print([cell.text for cell in row.cells])
